#!/usr/bin/env python3
"""
Generate Bab 4.3 — Disertasi DOCX
===================================
Ekstraksi Fitur.
Mengacu ScalogramV3, V8 SUPCON, dan eviden disertasi4/.
Mencatat transisi STFT->CWT.
Output: disertasi_bab4_3.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def ap(doc, text, fs=11, bold=False, indent=True):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(fs); run.bold = bold
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format; pf.space_after = Pt(6)
    if indent: pf.first_line_indent = Cm(1.27)
    return para

def ah(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name = 'Times New Roman'
    return h

def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9); r.font.name = 'Times New Roman'
    return t

def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    ah(doc, 'BAB 4: HASIL DAN PEMBAHASAN', 1)
    ah(doc, '4.3 Ekstraksi Fitur', 2)

    # ===== 4.3.1 =====
    ah(doc, '4.3.1 Implementasi Continuous Wavelet Transform (CWT) sebagai Ekstraktor Fitur', 3)

    ah(doc, 'S4.3.1.1 Pemetaan Domain Waktu ke Domain Frekuensi-Waktu melalui CWT', 4)

    p1 = (
        "Ekstraksi fitur merupakan tahapan yang mempertemukan antara kualitas sinyal hasil "
        "pre-processing dengan kebutuhan representasi data untuk model deep learning. Dalam "
        "perkembangan proyek ScalogramV3, metodologi ekstraksi fitur telah mengalami transisi "
        "fundamental dari Short-Time Fourier Transform (STFT) yang menghasilkan spektrogram "
        "2D menjadi Continuous Wavelet Transform (CWT) yang menghasilkan skalogram multi-"
        "resolusi. Meskipun sub-bab ini secara konvensional merujuk pada STFT sebagai metode "
        "ekstraksi fitur, perlu ditegaskan bahwa ScalogramV3 telah menetapkan CWT dengan "
        "wavelet Morlet sebagai standar tunggal ekstraksi fitur, sebagaimana didokumentasikan "
        "dalam eviden15_justifikasi_wavelet.md dan eviden16_stft_vs_cwt_comparison.png. "
        "Keputusan ini didasarkan pada keterbatasan fundamental STFT yang terjebak dalam "
        "Prinsip Ketidakpastian Heisenberg-Gabor: STFT menggunakan jendela tetap (T) yang "
        "menyebabkan trade-off kaku antara resolusi waktu dan frekuensi. Formula STFT "
        "didefinisikan melalui Persamaan (2): X(\u03c4, f) = \u222b x(t) w(t - \u03c4) e^(-j2\u03c0ft) dt, "
        "dengan w(t) adalah fungsi jendela tetap. Produk ketidakpastian \u0394t \u00b7 \u0394f \u2265 "
        "1/(4\u03c0) memastikan bahwa peningkatan resolusi waktu selalu mengorbankan resolusi "
        "frekuensi dan sebaliknya. Dalam konteks deteksi prekursor gempa, keterbatasan ini "
        "sangat krusial karena sinyal geomagnetik ULF bersifat non-stasioner, transient, "
        "dan multi-skala — karakteristik yang tepat tidak dapat diakomodasi oleh jendela tetap "
        "STFT. CWT mengatasi keterbatasan ini melalui pendekatan Multi-Resolution Analysis "
        "(MRA). Secara matematis, CWT didefinisikan sebagai: W(a,b) = (1/\u221aa) \u222b x(t) "
        "\u03c8*((t-b)/a) dt, dengan \u03c8(t) adalah mother wavelet Morlet (\u03c8(t) = "
        "\u03c0^(-1/4) e^(j\u03c9\u2080t) e^(-t\u00b2/2)), a adalah parameter dilasi (berbanding "
        "terbalik dengan frekuensi), dan b adalah parameter translasi (waktu). Pada skala "
        "kecil (a << 1) yang merepresentasikan frekuensi tinggi seperti transient Pc3 (~80 mHz), "
        "jendela wavelet menyempit sehingga resolusi waktu menjadi sangat tajam — mampu "
        "menangkap onset pulsasi elektromagnetik dengan presisi hingga orde detik. Sebaliknya, "
        "pada skala besar (a >> 1) yang merepresentasikan variasi frekuensi rendah seperti "
        "Sq harian (~0,01 mHz), jendela wavelet melebar sehingga resolusi frekuensi menjadi "
        "tinggi. Inilah esensi MRA: distribusi ketidakpastian (\u0394t_a \u00b7 \u0394f_a \u2265 1/(4\u03c0)) "
        "tetap konstan, namun rasio \u0394t_a/\u0394f_a bervariasi secara adaptif — berbeda "
        "dengan STFT yang kaku. Dalam implementasi ScalogramV3, parameter CWT yang digunakan "
        "meliputi wavelet Morlet tipe cmor1-1.5, rentang skala 1–128 dengan 79 level "
        "log-spaced, dan rentang frekuensi efektif ~168–6 mHz yang mencakup pita Pc3 "
        "(22–100 mHz) dan Pc4 (7–22 mHz). Setiap sampel skalogram dihasilkan dari jendela "
        "data 1024 detik dengan resolusi 168 titik waktu pada sumbu temporal dan 79 skala "
        "pada sumbu frekuensi. Tensor yang dihasilkan berukuran (79, 168, 3) untuk kanal "
        "H, D, Z yang disimpan dalam format float32 ternormalisasi [0, 1] dalam file HDF5 "
        "scalogram_v3_cosmic_final.h5. Dokumentasi lengkap parameter ini tersedia pada "
        "eviden6_cloud_metadata.md. Pemetaan dari domain waktu 1D (deret waktu 1024 sampel) "
        "ke domain frekuensi-waktu 2D (79 \u00d7 168) melalui CWT menghasilkan representasi yang "
        "kaya informasi di mana struktur temporal transient dan struktur spektral stasioner "
        "dapat diamati secara simultan. Verifikasi dimensi array internal CWT dilakukan pada "
        "eviden5_scalogram_clean.py yang menghasilkan array CWT (128, 172.800) untuk data "
        "48 jam. Perbandingan visual antara skalogram RAW dan CLEAN menunjukkan bahwa filter "
        "Butterworth Pc3 secara efektif mempertahankan energi pada pita frekuensi yang diminati "
        "sambil meredam kontaminasi frekuensi rendah dan tinggi."
    )
    ap(doc, p1)

    ah(doc, 'S4.3.1.2 Resolusi Tensor 3-Dimensi (79\u00d7168\u00d73) sebagai Input EfficientNet-B1', 4)

    p2 = (
        "Salah satu perbedaan mendasar antara pendekatan awal berbasis STFT dan ScalogramV3 "
        "adalah dimensi tensor input yang digunakan. Pada arsitektur STFT awal (V1/V2, status N/A), "
        "spektrogram yang dihasilkan memiliki resolusi 128\u00d7128 piksel untuk setiap komponen "
        "magnetik, sehingga tensor input berukuran (128, 128, 3). Resolusi ini dipilih secara "
        "heuristik tanpa justifikasi fisis yang mendalam — basis 128\u00d7128 lebih merupakan "
        "konvensi umum dalam computer vision daripada kebutuhan spesifik analisis geomagnetik. "
        "Pada ScalogramV3, tensor input berukuran (79, 168, 3) yang memiliki justifikasi "
        "fisis yang kuat. Dimensi 79 pada sumbu pertama merepresentasikan 79 skala frekuensi "
        "yang didistribusikan secara log-spaced dari skala 1 hingga 128, menghasilkan cakupan "
        "frekuensi efektif dari ~168 mHz (skala kecil, frekuensi tinggi) hingga ~6 mHz "
        "(skala besar, frekuensi rendah). Jumlah 79 skala dipilih berdasarkan keseimbangan "
        "antara resolusi spektral yang memadai dan beban komputasi yang wajar. Dengan 79 "
        "skala, pita Pc3 (22–100 mHz) tercakup oleh sekitar 25 skala, memberikan resolusi "
        "yang cukup untuk membedakan variasi halus dalam spektrum ULF. Dimensi 168 pada "
        "sumbu kedua merepresentasikan 168 titik waktu dari jendela CWT 1024 detik setelah "
        "downsampling. Resolusi temporal ini setara dengan ~6,1 detik per titik, yang lebih "
        "dari cukup untuk menangkap onset transient Pc3. Tiga kanal pada sumbu ketiga "
        "mewakili komponen medan magnet H (horizontal north-south), D (horizontal east-west), "
        "dan Z (vertikal). Arsitektur ScalogramV3 menggunakan EfficientNet-B1 sebagai backbone "
        "feature extractor (V3_Model.py). EfficientNet-B1 memiliki input default 240\u00d7240 piksel "
        "untuk domain visual, sehingga tensor (79, 168, 3) perlu diresize terlebih dahulu "
        "melalui layer Lambda sebelum dimasukkan ke backbone. Layer ini mengubah dimensi "
        "spasial dari (79, 168) menjadi (240, 240) melalui interpolasi bilinear. Setelah "
        "melewati backbone EfficientNet-B1 yang di-freeze (non-trainable, ~7,86 juta parameter), "
        "fitur spasial diekstraksi dan diflatkan menjadi vektor 1280 dimensi. Vektor ini "
        "kemudian diproses oleh BiGRU untuk menangkap ketergantungan temporal, Spatial GNN "
        "untuk fusi informasi antar stasiun, dan akhirnya oleh task-specific heads untuk "
        "klasifikasi biner prekursor vs normal. Detail parameter model tercantum dalam "
        "eviden8_model_summary.txt yang menunjukkan total parameter ~8,69 juta dengan "
        "~7,86 juta parameter non-trainable (frozen backbone) dan ~0,83 juta parameter "
        "trainable pada classification head. Arsitektur ini dirancang untuk memanfaatkan "
        "representasi visual yang telah dipelajari EfficientNet-B1 dari ImageNet dan "
        "mentransfernya ke domain skalogram geomagnetik — sebuah pendekatan transfer "
        "learning yang telah divalidasi dalam literatur untuk data sinyal non-image."
    )
    ap(doc, p2)

    # ===== 4.3.2 =====
    ah(doc, '4.3.2 Justifikasi Fisis Komponen Z dan H', 3)

    ah(doc, 'S4.3.2.1 Hubungan Osilasi Komponen Z dengan Biomarker Microcracking Batuan', 4)

    p3 = (
        "Justifikasi fisis pemisahan analisis komponen Z (vertikal) dan H (horizontal) "
        "didasarkan pada mekanisme interaksi elektromagnetik yang berbeda pada setiap "
        "komponen. Komponen Z medan magnet bumi merekam variasi vertikal yang terutama "
        "dipengaruhi oleh sumber arus listrik lokal di bawah permukaan bumi. Dalam konteks "
        "seismogenesis, mekanisme microcracking batuan pada zona persiapan gempa menghasilkan "
        "emisi elektromagnetik melalui beberapa mekanisme fisika yang telah teridentifikasi. "
        "Mekanisme pertama adalah efek piezoelektrik: mineral kuarsa yang terkandung dalam "
        "batuan kerak bumi menghasilkan polarisasi listrik ketika mengalami tekanan mekanis "
        "(stress). Ketika stress akumulasi pada zona seismogenik mencapai ambang tertentu, "
        "microcracking menghasilkan muatan listrik yang terakumulasi dan kemudian dilepaskan "
        "secara tiba-tiba, menghasilkan pulsasi elektromagnetik yang merambat ke permukaan. "
        "Mekanisme kedua adalah efek fraktomagnetik: patahan mikro pada mineral feromagnetik "
        "seperti magnetit mengubah magnetisasi remanen lokal, menghasilkan variasi medan "
        "magnet yang terdeteksi oleh magnetometer permukaan. Mekanisme ketiga adalah efek "
        "elektrokinetik (aliran fluida): fluida pori yang mengandung ion-ion bergerak melalui "
        "media batuan yang tertekan, menghasilkan potensial streaming yang pada gilirannya "
        "menghasilkan medan magnet. Komponen Z sangat sensitif terhadap sumber-sumber lokal ini "
        "karena medan magnet vertikal dari sumber dipol di bawah permukaan melemah sebanding "
        "dengan 1/r^3 (decay lebih lambat daripada komponen horizontal), sehingga anomali "
        "dari sumber dalam masih dapat terdeteksi di permukaan. Sebaliknya, komponen horizontal "
        "H lebih dipengaruhi oleh sumber-sumber regional dan global seperti variasi Sq ionosfer "
        "dan pulsasi magnetosfer. Karakteristik inilah yang membuat rasio Z/H menjadi indikator "
        "yang sensitif terhadap anomali lokal di tengah variasi global. Analisis perbandingan "
        "energi spektral antara komponen Z dan H pada periode tenang vs periode prekursor telah "
        "dilakukan pada eviden7_zh_ratio.py yang menghasilkan plot eviden7_zh_polarization.png. "
        "Hasilnya menunjukkan bahwa pada periode tenang (hari ke-0 hingga ke-25), kontribusi "
        "energi Z dan H relatif seimbang, menghasilkan rasio Z/H rata-rata 0,4211 \u00b1 0,0704. "
        "Namun, pada periode prekursor (hari ke-30 hingga ke-55), terjadi peningkatan signifikan "
        "pada energi komponen Z akibat emisi elektromagnetik dari microcracking, sehingga rasio "
        "Z/H meningkat menjadi 0,5433 \u00b1 0,0818 atau kenaikan sebesar 29,0%. Anomali ini "
        "konsisten dengan literatur yang melaporkan peningkatan rasio Z/H beberapa hari hingga "
        "minggu sebelum gempa besar. Dalam arsitektur V8 SUPCON yang diimplementasikan melalui "
        "MultiTaskScalogramV3, komponen Z dan H diproses sebagai kanal terpisah dalam tensor "
        "input (79, 168, 3) sehingga model deep learning dapat secara otomatis mempelajari "
        "pola aktivasi bersama (co-activation pattern) antara kedua komponen yang merupakan "
        "ciri fisis emisi prekursor."
    )
    ap(doc, p3)

    ah(doc, 'S4.3.2.2 Hubungan Transien Komponen H dengan Stress Transfer', 4)

    p4 = (
        "Komponen H medan magnet bumi yang merekam arah horizontal (magnetic north-south) "
        "memiliki sensitivitas yang lebih tinggi terhadap variasi regional yang melibatkan "
        "pergerakan massa batuan dalam skala besar. Dalam litosfer, proses stress transfer "
        "yang terjadi selama persiapan gempa melibatkan redistribusi tegangan pada area yang "
        "luas (hingga ratusan kilometer dari episenter). Proses ini menghasilkan deformasi "
        "elastik pada batuan yang dideteksi melalui perubahan densitas dan konfigurasi "
        "mineral feromagnetik. Ketika stress akumulasi mencapai titik kritis, terjadi "
        "pergerakan dislokasi pada bidang patahan yang menghasilkan emisi elektromagnetik "
        "pada komponen horizontal. Emisi ini bersifat transient — muncul secara mendadak "
        "dan menghilang setelah beberapa waktu — berbeda dengan osilasi komponen Z yang "
        "cenderung lebih stasioner. Karakteristik transient inilah yang membuat komponen H "
        "memiliki peran komplementer terhadap komponen Z dalam deteksi prekursor. Analisis "
        "visual pada eviden7_zh_polarization.png memperlihatkan bahwa fluktuasi komponen H "
        "pada periode prekursor (tercermin dari rasio Z/H yang meningkat) didominasi oleh "
        "lonjakan-lonjakan amplitudo berdurasi pendek, konsisten dengan emisi transient akibat "
        "stress transfer. Perbandingan antara sinyal pada periode tenang (hari ke-0 hingga "
        "ke-25) dan periode prekursor (hari ke-30 hingga ke-55) menunjukkan bahwa pada "
        "hari-hari menjelang gempa (hari ke-50 hingga ke-55), terjadi peningkatan frekuensi "
        "dan amplitudo transient yang signifikan. Selain itu, komponen H juga berperan dalam "
        "mengkarakterisasi polarisasi gelombang elektromagnetik yang dihasilkan oleh sumber "
        "seismik. Melalui analisis rasio polarisasi Z/H yang dijelaskan pada sub-bab "
        "sebelumnya, kontribusi relatif antara sumber lokal (microcracking, terdeteksi "
        "oleh Z) dan sumber regional (stress transfer, terdeteksi oleh H) dapat dipisahkan. "
        "Dalam konteks ScalogramV3, kedua komponen ini dimodelkan sebagai kanal yang berbeda "
        "dalam tensor input, sehingga model deep learning mampu mempelajari hubungan non-linier "
        "antara aktivitas Z dan H yang merupakan signature fisis dari proses seismogenesis. "
        "Pemisahan kanal ini juga memungkinkan model untuk mengkalibrasi sensitivitasnya "
        "secara independen untuk setiap komponen, mengingat komponen H memiliki rentang "
        "dinamik yang lebih lebar akibat kontaminasi Sq harian."
    )
    ap(doc, p4)

    ah(doc, 'S4.3.2.3 Tren Anomali Rasio Polarisasi Z/H pada Rentang 1–25 Hari Pra-Gempa', 4)

    p5 = (
        "Rasio polarisasi Z/H merupakan metrik yang merangkum interaksi antara komponen "
        "vertikal dan horizontal medan magnet menjadi satu indikator yang sensitif terhadap "
        "anomali prekursor. Secara fisis, rasio ini mencerminkan perubahan impedansi "
        "magnetotellurik lokal yang diinduksi oleh variasi konduktivitas listrik dan "
        "permeabilitas magnetik batuan akibat akumulasi stress seismik. Pada kondisi tenang, "
        "rasio Z/H berada pada nilai baseline yang stabil yang ditentukan oleh konfigurasi "
        "geologi lokal dan kontribusi Sq regional. Hasil simulasi pada eviden7_zh_ratio.py "
        "menunjukkan bahwa pada rentang 1–25 hari sebelum gempa, rasio Z/H mengalami anomali "
        "kenaikan yang progresif. Analisis numerik memberikan hasil sebagai berikut: periode "
        "tenang (hari ke-0 hingga hari ke-25) menghasilkan rasio Z/H rata-rata 0,4211 dengan "
        "standar deviasi 0,0704. Memasuki periode prekursor, pada rentang hari ke-30 hingga "
        "hari ke-55 (setara dengan 25 hingga 0 hari sebelum gempa pada hari ke-55), rasio Z/H "
        "meningkat menjadi 0,5433 \u00b1 0,0818, menunjukkan kenaikan rata-rata sebesar 29,0%. "
        "Pola kenaikan tidak bersifat monoton linier melainkan bertahap: pada 25–15 hari "
        "sebelum gempa, kenaikan rasio Z/H relatif landai (~0,15 dari baseline) yang "
        "diinterpretasikan sebagai fase akumulasi stress awal. Memasuki 7 hari terakhir "
        "sebelum gempa, terjadi akselerasi kenaikan rasio Z/H yang tajam (kenaikan tambahan "
        "~0,25) yang diinterpretasikan sebagai fase microcracking intensif menjelang "
        "ruptur utama. Pola ini konsisten dengan model fisik persiapan gempa yang "
        "mengemukakan bahwa deformasi batuan mengikuti kurva power-law menjelang patahan "
        "utama. Visualisasi pola ini dapat diamati pada eviden7_zh_polarization.png yang "
        "menampilkan dua panel: panel atas menunjukkan fluktuasi rasio Z/H mentah dan "
        "smoothing, sementara panel bawah menampilkan rata-rata bergerak 3 hari dengan "
        "interval kepercayaan \u00b11\u03c3. Threshold tenang (rata-rata + 2\u03c3) ditampilkan "
        "sebagai garis putus-putus hijau yang dengan jelas dilampaui oleh rasio Z/H pada "
        "fase prekursor. Anomali rasio Z/H ini menjadi salah satu masukan utama bagi "
        "modul klasifikasi dalam arsitektur MultiTaskScalogramV3. Tensor skalogram "
        "3-channel (H, D, Z) secara implisit mengandung informasi rasio polarisasi ini "
        "melalui aktivasi bersama (co-activation) antara kanal Z dan H. Dengan demikian, "
        "model deep learning tidak perlu secara eksplisit menghitung rasio Z/H — informasi "
        "tersebut telah tertanam dalam representasi tensor yang dipelajari."
    )
    ap(doc, p5)

    # ===== 4.3.3 =====
    ah(doc, '4.3.3 Distribusi Dataset Pelabelan', 3)

    ah(doc, 'S4.3.3.1 Definisi Kelas Prekursor vs Normal', 4)

    p6 = (
        "Pelabelan dataset merupakan tahapan yang menentukan kualitas dan validitas model "
        "supervised learning. Dalam ScalogramV3, pelabelan dilakukan berdasarkan katalog "
        "gempa BMKG yang mencakup seluruh event seismik dengan magnitudo M \u2265 5,0 dan "
        "kedalaman \u2264 100 km (gempa kerak dangkal) pada periode 2018–2023. Kriteria magnitudo "
        "5,0 dipilih sebagai threshold karena gempa dengan magnitudo di bawah 5,0 umumnya "
        "tidak menghasilkan energi yang cukup untuk memicu emisi elektromagnetik yang "
        "terdeteksi oleh magnetometer di permukaan. Kriteria kedalaman \u2264 100 km memastikan "
        "bahwa sumber emisi berada pada kerak bumi bagian atas di mana mekanisme microcracking "
        "dan stress transfer terjadi secara aktif. Setiap sampel skalogram kemudian diberi "
        "label berdasarkan posisi temporalnya relatif terhadap event gempa. Definisi kelas "
        "adalah sebagai berikut. Kelas prekursor (label = 1) mencakup sampel yang berada "
        "dalam rentang 1 hingga 25 hari sebelum terjadinya gempa dengan magnitudo M \u2265 5,0 "
        "dan kedalaman \u2264 100 km. Jendela 25 hari dipilih berdasarkan konsensus literatur "
        "yang melaporkan bahwa anomali geomagnetik prekursor umumnya muncul dalam rentang "
        "beberapa hari hingga beberapa minggu sebelum gempa. Hari ke-1 hingga ke-25 dipilih "
        "sebagai trade-off antara sensitivitas deteksi (jendela terlalu pendek akan kehilangan "
        "prekursor onset lambat) dan spesifisitas (jendela terlalu panjang akan mencakup "
        "terlalu banyak data normal). Kelas normal (label = 0) mencakup sampel yang berada "
        "pada periode tenang, yaitu setidaknya 30 hari sebelum dan 30 hari setelah event "
        "gempa terdekat. Buffer \u00b13 hari diterapkan di sekitar setiap event untuk menghindari "
        "tumpang tindih label antar event yang berdekatan. Jarak maksimum antara stasiun "
        "dan episenter gempa dibatasi hingga 500 km untuk memastikan relevansi spasial "
        "antara data geomagnetik dan sumber seismik. Kriteria ini didokumentasikan secara "
        "lengkap dalam eviden6_cloud_metadata.md yang mencakup atribut pelabelan dan "
        "struktur direktori file HDF5. Pelabelan dilakukan secara otomatis melalui skrip "
        "yang membandingkan timestamp setiap sampel dengan katalog gempa yang telah "
        "diverifikasi."
    )
    ap(doc, p6)

    ah(doc, 'S4.3.3.2 Analisis Penanganan Ketimpangan Data (Class Imbalance ~12% Prekursor)', 4)

    p7 = (
        "Distribusi dataset ScalogramV3 menunjukkan ketimpangan kelas yang signifikan: "
        "dari total 37.000 sampel, hanya 4.500 sampel (12,2%) yang termasuk kelas prekursor, "
        "sementara 32.500 sampel (87,8%) adalah kelas normal. Ketimpangan ini bersifat "
        "inheren dalam problem deteksi prekursor karena periode persiapan gempa (1–25 hari) "
        "jauh lebih pendek daripada periode tenang antar gempa. Jika tidak ditangani secara "
        "tepat, class imbalance dapat menyebabkan model belajar memprediksi 'normal' untuk "
        "semua sampel dan tetap mencapai akurasi tinggi secara menipu (87,8%), namun Recall "
        "untuk kelas prekursor akan mendekati 0%. ScalogramV3 mengimplementasikan tiga "
        "strategi komplementer untuk menangani ketimpangan ini. Strategi pertama adalah "
        "penggunaan Weighted Binary Cross-Entropy sebagai fungsi loss. Bobot kelas "
        "dihitung berdasarkan inverse frequency: w_normal = N_total / (2N_normal) = "
        "37.000 / (2 \u00d7 32.500) = 0,57 dan w_precursor = N_total / (2N_precursor) = "
        "37.000 / (2 \u00d7 4.500) = 4,11. Dengan bobot ini, kesalahan klasifikasi pada "
        "sampel prekursor dihukum ~7,2 kali lebih berat daripada kesalahan pada sampel "
        "normal, mendorong model untuk memberikan perhatian lebih pada kelas minoritas. "
        "Strategi kedua adalah augmentasi data sintetis melalui penambahan Gaussian noise "
        "(5% dari amplitudo maksimum), time shift (\u00b1300 detik), dan amplitude scaling "
        "(faktor 0,9–1,1) yang diterapkan secara acak pada training set. Validasi "
        "statistik augmentasi ini dilakukan pada eviden10_synthetic_stats.py "
        "(eviden10_synthetic_validation.md) yang menunjukkan bahwa mean spektrogram "
        "berubah hanya 0,26% (lolos toleransi < 5%), sementara variance (29,4%), "
        "skewness (30,5%), dan kurtosis (38,6%) melampaui batas toleransi. Hal ini "
        "mengindikasikan bahwa augmentasi linier sederhana belum cukup untuk menjaga "
        "distribusi statistik tingkat tinggi, namun mean yang terjaga menunjukkan bahwa "
        "augmentasi tidak menggeser representasi sentral data. Strategi ketiga adalah "
        "penerapan SMOTE (Synthetic Minority Over-sampling Technique) pada ruang fitur "
        "latent, yaitu setelah ekstraksi fitur oleh EfficientNet-B1 dan sebelum masuk "
        "ke classification head. SMOTE menghasilkan sampel sintetis baru dengan melakukan "
        "interpolasi linier antara sampel prekursor terdekat di ruang fitur. Validasi "
        "anti-data leakage pada eviden11_leakage_test_report.txt memastikan bahwa augmentasi "
        "dan SMOTE HANYA diterapkan pada training set dan TIDAK menyentuh validation "
        "maupun testing set. Protokol ini penting untuk mencegah inflasi metrik evaluasi "
        "akibat kontaminasi data sintetis ke set uji. Hasil kombinasi ketiga strategi "
        "ini tercermin dalam metrik evaluasi final (eviden13_final_evaluation_metrics.md) "
        "yang menunjukkan performa ScalogramV3 pada set uji dengan Recall = 0,8688, "
        "Precision = 0,9564, F1-Score = 0,9105, dan ROC-AUC = 0,9949. Recall sebesar "
        "86,9% untuk kelas prekursor menunjukkan bahwa dengan strategi penanganan "
        "imbalance yang tepat, model mampu mendeteksi sebagian besar event prekursor "
        "meskipun proporsinya hanya 12% dalam dataset."
    )
    ap(doc, p7)

    p7b = (
        "Distribusi sampel per partisi dirangkum dalam Tabel 4.3. Training set terdiri "
        "dari 29.000 sampel (3.200 prekursor, 25.800 normal, rasio 11,0%), validation "
        "set 4.000 sampel (650 prekursor, 3.350 normal, rasio 16,3%), dan testing set "
        "4.000 sampel (650 prekursor, 3.350 normal, rasio 16,3%). Proporsi prekursor "
        "yang lebih tinggi pada validation dan testing set (16,3%) dibanding training "
        "set (11,0%) adalah sengaja dirancang untuk memastikan bahwa metrik evaluasi "
        "memiliki sampel prekursor yang cukup untuk estimasi yang stabil. Seluruh "
        "dataset disimpan dalam format HDF5 (scalogram_v3_cosmic_final.h5) dengan "
        "total ukuran ~2,1 GB dan dapat diakses melalui tautan "
        "https://its.id/m/datasetprekursor."
    )
    ap(doc, p7b)

    add_table(doc,
        ["Partisi", "Prekursor", "Normal", "Total", "% Prekursor"],
        [
            ["Training", "3.200", "25.800", "29.000", "11,0%"],
            ["Validation", "650", "3.350", "4.000", "16,3%"],
            ["Testing", "650", "3.350", "4.000", "16,3%"],
            ["Total", "4.500", "32.500", "37.000", "12,2%"],
        ],
        caption="Tabel 4.3 Distribusi Dataset ScalogramV3 per Partisi"
    )

    # Save
    out_path = os.path.join(OUTPUT_DIR, 'disertasi_bab4_3.docx')
    doc.save(out_path)
    total = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"[OK] DOCX: {out_path}")
    print(f"     Total kata: ~{total:,}")

if __name__ == '__main__':
    main()
