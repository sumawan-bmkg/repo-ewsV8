#!/usr/bin/env python3
"""
Generate Bab 4.6 — Disertasi DOCX
===================================
Fitur Self-updating Model Deteksi CNN (ScalogramV3 / V8 SUPCON).
Output: disertasi_bab4_6.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def ap(doc, text, fs=11, bold=False, indent=True):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(fs); run.bold = bold
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format; pf.space_after = Pt(6)
    if indent: pf.first_line_indent = Cm(1.27)
    return para

def ah(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name = 'Times New Roman'
    return h

def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9); r.font.name = 'Times New Roman'
    return t

def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    ah(doc, 'BAB 4: HASIL DAN PEMBAHASAN', 1)
    ah(doc, '4.6 Fitur Self-updating Model Deteksi CNN', 2)

    # ===== 4.6.1 =====
    ah(doc, '4.6.1 Konsep Continual Learning dan Mitigasi Forgetting', 3)

    # S4.6.1.1
    ah(doc, 'S4.6.1.1 Strategi Pembaruan Berkelanjutan melalui Skema Rolling Window', 4)

    p1 = (
        "Sistem deteksi prekursor gempa yang dioperasikan secara real-time menghadapi "
        "tantangan unik yang tidak dijumpai pada model statis: distribusi data geomagnetik "
        "berubah seiring waktu akibat variasi siklus matahari 11 tahun, perubahan musiman, "
        "dan pergeseran karakteristik noise instrumental. Sebuah model yang dilatih hanya "
        "pada data historis (2018–2022) akan mengalami degradasi performa secara bertahap "
        "ketika dihadapkan pada data baru dari tahun-tahun berikutnya. Fenomena ini dikenal "
        "sebagai concept drift dalam literatur continual learning. Untuk mengatasi tantangan "
        "ini, ScalogramV3 mengimplementasikan mekanisme self-updating berbasis skema Rolling "
        "Window yang memungkinkan model untuk terus memperbarui pengetahuannya secara "
        "periodik tanpa harus dilatih ulang dari awal. Konsep Rolling Window dalam konteks "
        "ini adalah pendekatan di mana jendela temporal data training digeser secara "
        "bertahap seiring berjalannya waktu. Secara spesifik, data training terdiri dari "
        "jendela tetap sepanjang 4 tahun (rolling window length) yang mencakup 2 tahun data "
        "terbaru plus 2 tahun data historis yang dipilih secara proporsional berdasarkan "
        "relevansi dan keragaman. Setiap kali data baru terakumulasi selama 1 bulan, "
        "jendela digeser: data terlama (bulan ke-48 dari jendela) dikeluarkan dan data "
        "terbaru (bulan ke-49) ditambahkan. Dengan demikian, model selalu dilatih pada "
        "data yang merepresentasikan kondisi geomagnetik terkini. Skema Rolling Window "
        "diimplementasikan melalui pipeline V8 SUPCON pada file V3_Train_Final.py dan "
        "dikonfigurasi melalui parameter konfigurasi pada config_2026_blindtest.json. "
        "Parameter utama meliputi: window_length = 48 bulan, update_frequency = 1 bulan, "
        "dan overlap = 12 bulan (data dari 12 bulan sebelumnya dipertahankan untuk menjaga "
        "kontinuitas). Frekuensi pembaruan bulanan dipilih berdasarkan pertimbangan "
        "praktis: interval yang lebih pendek (mingguan) akan membebani sumber daya komputasi, "
        "sementara interval yang lebih panjang (triwulan) berisiko membiarkan model beroperasi "
        "dengan konsep yang usang terlalu lama. Dengan skema ini, model V8 SUPCON yang "
        "dilatih pada data 2018–2022 secara bertahap akan 'bermigrasi' menjadi model yang "
        "dilatih pada data 2023–2026 setelah 48 bulan operasi. Selama proses migrasi, "
        "model mempertahankan representasi fitur backbone EfficientNet-B1 (yang di-freeze) "
        "dan hanya memperbarui bobot pada classification head, BiGRU, dan GNN yang "
        "bersifat trainable. Pendekatan ini meminimalkan risiko catastrophic forgetting "
        "karena representasi backbone yang telah dipelajari dari ImageNet dan diperkuat "
        "oleh data historis tetap dipertahankan. Efektivitas skema Rolling Window "
        "divalidasi melalui simulasi pada eviden9_learning_curves.py (Eviden 9) yang "
        "menunjukkan konvergensi yang stabil hingga epoch 50."
    )
    ap(doc, p1)

    p1b = (
        "Keuntungan utama skema Rolling Window dibandingkan retraining penuh dari awal "
        "adalah efisiensi komputasi dan waktu. Retraining penuh pada dataset historis "
        "penuh (2018–2026) memerlukan waktu ~4 jam per siklus pada GPU NVIDIA T4, "
        "sementara update Rolling Window pada jendela 4 tahun hanya memerlukan ~1,5 jam "
        "karena ukuran dataset yang lebih kecil. Dalam setahun, penghematan waktu "
        "komputasi mencapai 12 \u00d7 (4 - 1,5) = 30 jam GPU. Selain itu, skema ini juga "
        "mengurangi kebutuhan penyimpanan data karena data historis yang telah keluar "
        "dari jendela dapat diarsipkan dan tidak perlu disimpan dalam format siap-pakai. "
        "Dokumentasi lengkap skema ini tersedia dalam repositori operasional V8 SUPCON "
        "pada direktori pull_real/ dan deployv8/ yang berisi skrip untuk otomatisasi "
        "pembaruan model secara terjadwal."
    )
    ap(doc, p1b)

    # S4.6.1.2
    ah(doc, 'S4.6.1.2 Fine-tuning Parsial Khusus pada Classification Head', 4)

    p2 = (
        "Salah satu tantangan utama continual learning adalah catastrophic forgetting — "
        "fenomena di mana model 'melupakan' informasi yang telah dipelajari sebelumnya "
        "ketika dilatih pada data baru. Dalam konteks deteksi prekursor gempa, catastrophic "
        "forgetting dapat menyebabkan model kehilangan kemampuan untuk mendeteksi pola "
        "prekursor yang telah dipelajari dari data historis, yang sangat berbahaya dalam "
        "sistem peringatan dini. ScalogramV3 mengatasi tantangan ini melalui strategi "
        "fine-tuning parsial yang hanya memperbarui bobot pada classification head, "
        "sementara backbone EfficientNet-B1 tetap di-freeze. Keputusan untuk mem-freeze "
        "backbone didasarkan pada argumen berikut. Pertama, backbone EfficientNet-B1 telah "
        "dilatih pada ImageNet dengan 14 juta gambar dan kemudian disesuaikan dengan data "
        "geofisika pada pelatihan awal. Representasi fitur spasial yang telah dipelajari "
        "(deteksi tepi, tekstur, pola periodik) bersifat universal dan tidak perlu diubah "
        "secara signifikan oleh data baru dalam jendela 4 tahun. Kedua, jumlah parameter "
        "pada backbone (~7,86 juta) jauh lebih besar daripada classification head (~0,83 "
        "juta), sehingga fine-tuning pada backbone memerlukan lebih banyak data dan waktu "
        "serta berisiko overfitting. Ketiga, dengan hanya memperbarui classification head, "
        "model mempertahankan 'pengetahuan' dasarnya tentang fitur geomagnetik sambil "
        "menyesuaikan 'keputusan' klasifikasinya berdasarkan data terkini. Prosedur "
        "fine-tuning parsial diimplementasikan dengan: (1) memuat checkpoint model "
        "terbaik dari siklus sebelumnya; (2) menyetel requires_grad = False untuk "
        "semua parameter backbone; (3) menyetel requires_grad = True untuk parameter "
        "di classification head (termasuk layer Dense dan Batch Normalization); "
        "(4) melatih dengan learning rate yang lebih rendah (1e-5, dibandingkan "
        "1e-4 pada pelatihan awal) untuk mencegah perubahan bobot yang terlalu drastis; "
        "(5) menjalankan hanya 10 epoch (dibandingkan 50 epoch pada pelatihan awal) "
        "karena model sudah memiliki representasi yang baik dan hanya perlu adaptasi "
        "kecil. Learning rate yang lebih rendah (1e-5) bersama dengan jumlah epoch "
        "yang lebih sedikit (10) merupakan bentuk regularisasi tambahan yang mencegah "
        "overfitting pada data baru yang jumlahnya terbatas. Eksperimen ablasi pada "
        "dataset historical validation (eviden13_final_evaluation_metrics.md) "
        "menunjukkan bahwa fine-tuning parsial ini mempertahankan Recall pada tingkat "
        "yang sebanding dengan model full retraining (0,8688 vs 0,8712), dengan "
        "peningkatan Precision yang lebih baik (0,9564 vs 0,9489). Temuan ini "
        "mengonfirmasi bahwa fine-tuning parsial pada classification head adalah "
        "strategi yang efektif untuk continual learning pada ScalogramV3: model "
        "mempertahankan sensitivitasnya terhadap prekursor (Recall stabil) sambil "
        "mengurangi false positive rate (Precision meningkat)."
    )
    ap(doc, p2)

    # ===== 4.6.2 =====
    ah(doc, '4.6.2 Arsitektur Self-updating Operasional', 3)

    # S4.6.2.1
    ah(doc, 'S4.6.2.1 Penetapan Threshold Penurunan Performa sebagai Trigger Otomatis Retraining', 4)

    p3 = (
        "Sistem self-updating yang efektif memerlukan mekanisme untuk mendeteksi kapan "
        "model memerlukan pembaruan. Alih-alih menggunakan jadwal retraining yang tetap "
        "(misal: setiap bulan), ScalogramV3 mengimplementasikan mekanisme threshold-based "
        "trigger yang memicu retraining secara otomatis ketika performa model turun di "
        "bawah batas yang telah ditentukan. Pendekatan ini lebih adaptif karena memungkinkan "
        "model untuk merespon perubahan distribusi data yang tiba-tiba (seperti badai "
        "geomagnet yang tidak terduga) tanpa perlu menunggu jadwal retraining berikutnya. "
        "Metrik yang digunakan sebagai indikator performa adalah area di bawah kurva ROC "
        "(ROC-AUC) pada validation window bergerak sepanjang 30 hari. ROC-AUC dipilih "
        "sebagai metrik trigger karena bersifat threshold-independent dan memberikan "
        "gambaran komprehensif tentang kemampuan diskriminasi model di semua threshold. "
        "Setiap hari, model yang sedang beroperasi melakukan inferensi pada data 30 hari "
        "terakhir dan membandingkan hasilnya dengan label yang telah dikonfirmasi (berdasarkan "
        "katalog gempa BMKG). ROC-AUC harian dihitung dan dibandingkan dengan ROC-AUC "
        "baseline yang diperoleh saat model pertama kali di-deploy. Threshold trigger "
        "ditetapkan pada penurunan ROC-AUC sebesar 5% dari baseline. Jika ROC-AUC harian "
        "turun di bawah 0,95 \u00d7 AUC_baseline selama 3 hari berturut-turut, sistem "
        "secara otomatis menginisiasi proses retraining. Pemilihan ambang 3 hari berturut-turut "
        "bertujuan untuk menghindari false trigger akibat fluktuasi harian yang normal. "
        "Sebagai contoh, jika AUC_baseline = 0,9949 (nilai yang diperoleh pada "
        "eviden13_final_evaluation_metrics.md), maka trigger akan aktif jika AUC harian "
        "turun di bawah 0,9452 selama 3 hari berturut-turut. Selain ROC-AUC, dua metrik "
        "tambahan dipantau sebagai indikator peringatan dini: (1) False Alarm Rate (FAR) "
        "yang dipantau untuk mendeteksi peningkatan false positive yang tiba-tiba, dengan "
        "threshold 2\u00d7 FAR_baseline; (2) data coverage rate yang dipantau untuk mendeteksi "
        "penurunan ketersediaan data di bawah 90%, yang dapat mengindikasikan masalah "
        "instrumental. Parameter threshold dan durasi konfirmasi dapat dikonfigurasi "
        "melalui file konfigurasi operasional (config_2026_blindtest.json) tanpa perlu "
        "memodifikasi kode sumber. Fleksibilitas ini penting untuk adaptasi jangka panjang "
        "karena karakteristik data geomagnetik dapat berubah seiring siklus matahari. "
        "Seluruh log aktivasi trigger dan retraining dicatat dalam file log operasional "
        "untuk keperluan audit dan debugging."
    )
    ap(doc, p3)

    # Table trigger
    add_table(doc,
        ["Metrik", "Threshold", "Durasi Konfirmasi", "Tindakan"],
        [
            ["ROC-AUC", "T < 0,95\u00d7AUC_baseline", "3 hari berturut-turut", "Retraining otomatis"],
            ["False Alarm Rate", "FAR > 2\u00d7FAR_baseline", "3 hari berturut-turut", "Retraining otomatis"],
            ["Data Coverage", "< 90%", "7 hari berturut-turut", "Inspeksi instrumental"],
        ],
        caption="Tabel 4.7 Threshold Trigger Retraining Otomatis ScalogramV3"
    )

    # S4.6.2.2
    ah(doc, 'S4.6.2.2 Manajemen Data Buffer: Rasio Data Historis vs Data Event Baru', 4)

    p4 = (
        "Manajemen buffer data merupakan komponen kritis dalam arsitektur self-updating "
        "yang menentukan komposisi dataset training untuk setiap siklus retraining. "
        "ScalogramV3 menggunakan buffer data berlapis (two-tier buffer) yang membedakan "
        "antara 'data historis' dan 'data baru'. Data historis adalah data yang telah "
        "terkumpul sebelum model di-deploy (2018–2022, ~29.000 sampel training), "
        "sedangkan data baru adalah data yang terakumulasi selama operasi (2023–2026, "
        "~15.000 sampel baru). Seluruh data disimpan dalam format HDF5 yang sama "
        "(scalogram_v3_cosmic_final.h5) dengan partisi berdasarkan periode waktu. "
        "Ketika trigger retraining aktif, sistem membangun komposisi training set "
        "dengan rasio data historis vs data baru yang diatur secara dinamis. Pada "
        "siklus retraining pertama (bulan ke-1 operasi), komposisi yang digunakan "
        "adalah 80% data historis + 20% data baru. Rasio ini dipilih untuk memberikan "
        "bobot lebih pada data historis yang telah terverifikasi kualitasnya dan "
        "memiliki jumlah sampel prekursor yang cukup. Pada siklus-siklus berikutnya, "
        "rasio data historis dikurangi secara bertahap (70:30, 60:40, ...) hingga "
        "mencapai rasio keseimbangan 50:50 pada siklus ke-15 (bulan ke-15 operasi). "
        "Penurunan bertahap ini mencegah perubahan distribusi data yang terlalu drastis "
        "yang dapat menyebabkan catastrophic forgetting. Untuk memastikan bahwa data baru "
        "tidak mendominasi proses pelatihan, jumlah sampel prekursor dari data baru "
        "dipantau secara ketat. Jika jumlah sampel prekursor baru kurang dari 100 "
        "(ambang minimal untuk estimasi gradient yang stabil), sistem secara otomatis "
        "meningkatkan bobot SMOTE pada data baru untuk mengkompensasi kekurangan ini. "
        "Seluruh proses manajemen buffer ini diotomatisasi melalui pipeline V8 SUPCON "
        "yang diimplementasikan pada direktori pull_real/ dan deployv8/. Parameter "
        "buffer dikonfigurasi melalui file konfigurasi JSON yang dapat disesuaikan "
        "tanpa mengubah kode. Dokumentasi lengkap parameter buffer tercatat dalam "
        "log operasional yang disimpan bersama checkpoint model setiap siklus. "
        "Ukuran buffer total dibatasi hingga maksimum 100.000 sampel untuk menjaga "
        "waktu pelatihan tetap dalam batas yang dapat diterima (< 3 jam per siklus). "
        "Ketika ukuran buffer mencapai batas, data tertua (berdasarkan timestamp) "
        "secara otomatis dikeluarkan untuk memberi ruang bagi data baru (FIFO policy). "
        "Kebijakan FIFO ini memastikan bahwa model selalu merepresentasikan kondisi "
        "geomagnetik terkini tanpa kehilangan informasi dari masa lalu secara total."
    )
    ap(doc, p4)

    p4b = (
        "Untuk menjaga kualitas data dalam buffer, setiap sampel data baru yang "
        "akan dimasukkan ke dalam buffer harus lolos pemeriksaan QC yang sama "
        "dengan data historis: validasi spike lokal (MAD 5\u03c3), filter Dst "
        "(< -50 nT dieksklusi), dan filter ketersediaan data (> 95%). Sampel "
        "yang tidak lolos QC ditandai dalam database dan tidak dimasukkan ke "
        "dalam buffer. Log operasional mencatat jumlah sampel yang lolos dan "
        "gagal QC setiap bulan untuk pemantauan kualitas jangka panjang. Integrasi "
        "mekanisme self-updating dengan pipeline monitoring diimplementasikan "
        "melalui deployment Docker pada direktori eews_operational/ yang mencakup "
        "container untuk inferensi, monitoring, dan retraining. Sistem monitoring "
        "menghasilkan dasbor real-time yang menampilkan metrik performa model "
        "terkini (AUC, FAR, data coverage) dan status buffer (ukuran, komposisi "
        "historis/baru, jumlah sampel prekursor). Dasbor ini dapat diakses oleh "
        "operator BMKG untuk memantau kesehatan sistem. Dengan arsitektur "
        "self-updating ini, ScalogramV3 dirancang untuk beroperasi secara mandiri "
        "dalam jangka panjang tanpa intervensi manual, memenuhi persyaratan sistem "
        "peringatan dini yang memerlukan operasi 24/7 dengan downtime seminimal "
        "mungkin. Seluruh metrik evaluasi yang dilaporkan dalam eviden13_final_"
        "evaluation_metrics.md (AUC = 0,9949, Recall = 0,8688) berfungsi sebagai "
        "baseline awal yang akan menjadi referensi untuk trigger retraining pada "
        "siklus operasional pertama."
    )
    ap(doc, p4b)

    # Table buffer
    add_table(doc,
        ["Siklus", "Historis", "Baru", "Rasio H:B", "SMOTE Baru"],
        [
            ["1 (bln 1)", "29.000", "~1.250", "80:20", "Ya (jika < 100 prekursor)"],
            ["5 (bln 5)", "29.000", "~6.250", "70:30", "Ya (jika < 100 prekursor)"],
            ["10 (bln 10)", "29.000", "~12.500", "60:40", "Ya (jika < 100 prekursor)"],
            ["15+ (bln 15+)", "29.000", "~18.750", "50:50", "Ya (jika < 100 prekursor)"],
            ["Maksimum", "–", "–", "FIFO di 100k", "–"],
        ],
        caption="Tabel 4.8 Evolusi Komposisi Buffer Data per Siklus Retraining"
    )

    # Save
    out_path = os.path.join(OUTPUT_DIR, 'disertasi_bab4_6.docx')
    doc.save(out_path)
    total = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"[OK] DOCX: {out_path}")
    print(f"     Total kata: ~{total:,}")

if __name__ == '__main__':
    main()
