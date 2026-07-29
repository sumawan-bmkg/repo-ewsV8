#!/usr/bin/env python3
"""
Merge all Bab 4 DOCX files into one file: bab4.docx
"""
import os
from docx import Document

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

FILES = [
    'disertasi_bab4_1.docx',
    'disertasi_bab4_2.docx',
    'disertasi_bab4_3.docx',
    'disertasi_bab4_4.docx',
    'disertasi_bab4_5.docx',
    'disertasi_bab4_6.docx',
    'disertasi_bab4_7.docx',
]

def merge_docx(input_paths, output_path):
    merged = Document()
    # Copy default style
    merged.styles['Normal'].font.name = 'Times New Roman'
    merged.styles['Normal'].font.size = Pt(12)

    first = True
    for path in input_paths:
        doc = Document(path)
        for element in doc.element.body:
            if first:
                merged.element.body.append(element)
            else:
                merged.element.body.append(element)
        first = False

    # Actually, python-docx doesn't support direct element appending well.
    # Better approach: iterate paragraphs and tables.
    merged.save(output_path)

# The proper way: use paragraph-by-paragraph + table merging
def merge_docx_proper(input_paths, output_path):
    merged = Document()
    merged.styles['Normal'].font.name = 'Times New Roman'
    merged.styles['Normal'].font.size = Pt(12)

    # Remove the default empty paragraph
    if merged.paragraphs:
        p = merged.paragraphs[0]._element
        p.getparent().remove(p)

    for i, path in enumerate(input_paths):
        doc = Document(path)
        for para in doc.paragraphs:
            new_para = merged.add_paragraph()
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            new_para.alignment = para.alignment
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
        for table in doc.tables:
            new_table = merged.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = table.style
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    new_cell = new_table.rows[r_idx].cells[c_idx]
                    new_cell.text = cell.text
                    for p in new_cell.paragraphs:
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)

    merged.save(output_path)
    print(f"[OK] Merged {len(input_paths)} files -> {output_path}")
    total_words = sum(len(p.text.split()) for p in merged.paragraphs)
    print(f"     Total kata: ~{total_words:,}")

if __name__ == '__main__':
    from docx.shared import Pt
    input_files = [os.path.join(OUTPUT_DIR, f) for f in FILES]
    output_file = os.path.join(OUTPUT_DIR, 'bab4.docx')

    missing = [f for f in input_files if not os.path.exists(f)]
    if missing:
        print(f"ERROR: Files missing: {missing}")
    else:
        merge_docx_proper(input_files, output_file)
